"""
storage/text_extractor.py — Extract plain text (and structure) from uploaded
documents.

Supported types: PDF (.pdf), Word (.docx), Markdown (.md), and plain text
(.txt). PDF parsing uses pdfplumber; DOCX uses python-docx. Each is imported
lazily so a missing optional dependency only fails when that file type is
actually uploaded.

extract_text() returns a flat string (kept for backward compatibility).
extract_sections() returns the same content split into heading-delimited
sections, which storage/chunker.py uses for structure-aware chunking instead
of pure character-count splitting.
"""

import logging
import re
from pathlib import Path
from typing import TypedDict

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}


class Section(TypedDict):
    heading: str | None
    text: str


def extract_text(path: str | Path) -> str:
    """
    Return the plain-text content of a document.

    Raises ValueError for unsupported file types.
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".md", ".txt"):
        return path.read_text(encoding="utf-8", errors="replace")

    raise ValueError(
        f"Unsupported document type {ext!r}. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
    )


def extract_sections(path: str | Path) -> list[Section]:
    """
    Return the document split into heading-delimited sections.

    DOCX uses the real paragraph style ("Heading 1", "Title", ...) — the most
    reliable structural signal available. Markdown uses '#' lines. PDF and
    plain TXT have no structural metadata, so a generic heuristic splits on
    short, unpunctuated lines that look like headings (e.g. "Security Policy"
    followed by paragraph text). Every section always has non-empty text;
    content before the first detected heading gets heading=None.
    """
    path = Path(path)
    ext = path.suffix.lower()

    if ext == ".docx":
        return _extract_docx_sections(path)
    if ext == ".md":
        return _split_markdown_sections(path.read_text(encoding="utf-8", errors="replace"))

    # PDF and plain TXT: extract flat text, then apply the heading heuristic.
    text = extract_text(path)
    return _split_heuristic_sections(text)


def _extract_pdf(path: Path) -> str:
    import pdfplumber

    parts: list[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                parts.append(page_text)
    text = "\n".join(parts)
    logger.info("[Extractor] PDF %s → %d chars", path.name, len(text))
    return text


def _extract_docx(path: Path) -> str:
    import docx  # python-docx

    document = docx.Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    logger.info("[Extractor] DOCX %s → %d chars", path.name, len(text))
    return text


def _extract_docx_sections(path: Path) -> list[Section]:
    import docx  # python-docx

    document = docx.Document(str(path))
    sections: list[Section] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        body = "\n".join(current_lines).strip()
        if body:
            sections.append({"heading": current_heading, "text": body})

    for p in document.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style_name = (p.style.name if p.style else "") or ""
        is_heading = style_name.startswith("Heading") or style_name == "Title"
        if is_heading:
            flush()
            current_heading = text
            current_lines = []
        else:
            current_lines.append(text)
    flush()

    logger.info("[Extractor] DOCX %s → %d sections", path.name, len(sections))
    return sections or [{"heading": None, "text": _extract_docx(path)}]


def _split_markdown_sections(text: str) -> list[Section]:
    sections: list[Section] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        body = "\n".join(current_lines).strip()
        if body:
            sections.append({"heading": current_heading, "text": body})

    for line in text.splitlines():
        m = re.match(r"^#{1,6}\s+(.+)$", line.strip())
        if m:
            flush()
            current_heading = m.group(1).strip()
            current_lines = []
        else:
            current_lines.append(line)
    flush()

    return sections or [{"heading": None, "text": text.strip()}]


# A line is treated as a heading when it's short, has no terminal sentence
# punctuation, and isn't itself blank — true of "Security Policy" but not of
# "Badges must be visible at all times." This is a heuristic (PDF/TXT carry no
# structural metadata) but works well for the policy/report-style documents
# this system targets.
_HEADING_LINE_RE = re.compile(r"^[^.!?;:]{1,80}$")


def _split_heuristic_sections(text: str) -> list[Section]:
    lines = [ln.strip() for ln in text.splitlines()]
    sections: list[Section] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def flush() -> None:
        body = "\n".join(current_lines).strip()
        if body:
            sections.append({"heading": current_heading, "text": body})

    for line in lines:
        if not line:
            continue
        looks_like_heading = bool(_HEADING_LINE_RE.match(line)) and len(line.split()) <= 10
        if looks_like_heading and current_lines:
            # Only treat as a new heading once we already have body text under
            # the current one — avoids a doc's very first line always being a
            # spurious "heading" with nothing else attached.
            flush()
            current_heading = line
            current_lines = []
        elif looks_like_heading and not current_lines and current_heading is None:
            current_heading = line
        else:
            current_lines.append(line)
    flush()

    return sections or [{"heading": None, "text": text.strip()}]
